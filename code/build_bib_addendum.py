"""
Build an addendum docx for Guy Miron's review covering items added today,
in the same two-column Citation | Notes layout as the 2026-06-08 master doc.
"""
import json, urllib.request, os, datetime
from collections import defaultdict
from docx import Document
from docx.shared import Inches

GROUP = '6579255'
API = 'https://api.zotero.org/groups/' + GROUP
HEAD = {'Zotero-API-Version': '3', 'Zotero-API-Key': os.environ['ZOTERO_API_KEY']}

def get(path, params=''):
    req = urllib.request.Request(API + path + ('?' + params if params else ''), headers=HEAD)
    return json.loads(urllib.request.urlopen(req).read())

# Collections + paths
colls_raw, start = [], 0
while True:
    page = get('/collections', f'limit=100&start={start}')
    if not page: break
    colls_raw += page
    start += 100
colls = {c['key']: c['data'] for c in colls_raw}

def coll_path(k):
    parts = []
    while k:
        d = colls.get(k)
        if not d: break
        parts.insert(0, d['name'])
        k = d.get('parentCollection')
    return ' > '.join(parts)

# Items added today
today = datetime.date.today().isoformat()
items_today, start = [], 0
while True:
    page = get('/items', f'limit=100&start={start}&sort=dateAdded&direction=desc')
    if not page: break
    keep = [it for it in page if it['data'].get('dateAdded','').startswith(today)
            and it['data'].get('itemType') not in ('attachment','note')]
    items_today += keep
    if all(it['data'].get('dateAdded','') < today for it in page) and start > 0:
        break
    start += 100

P9 = '63X662X9'
review_items   = [it for it in items_today if P9 not in it['data'].get('collections', [])]
registry_items = [it for it in items_today if P9 in it['data'].get('collections', [])]
print(f"Review items: {len(review_items)}; P9 registry: {len(registry_items)}")

by_coll = defaultdict(list)
for it in review_items:
    for ck in it['data'].get('collections', []):
        by_coll[ck].append(it)

def cite(d):
    auths = '; '.join(
        f"{c.get('lastName','')}, {c.get('firstName','')}".strip(', ')
        if c.get('creatorType') in ('author','editor','interviewee')
        else c.get('name','')
        for c in d.get('creators', [])
    ) or '[no author]'
    year = (d.get('date','') or '')[:4]
    title = d.get('title','').rstrip('.')
    container = d.get('publicationTitle') or d.get('bookTitle') or d.get('proceedingsTitle') or ''
    publisher = d.get('publisher','')
    place = d.get('place','')
    pages = d.get('pages','')
    url = d.get('url','')
    s = f"{auths} ({year}). {title}."
    if container: s += f" In {container}"
    if pages:     s += f", S. {pages}"
    if publisher: s += f". {place + ': ' if place else ''}{publisher}"
    s = s.rstrip('.') + '.'
    if url: s += f" {url}"
    return s

def tag_set(d):
    return {t.get('tag','') if isinstance(t,dict) else t for t in d.get('tags', []) or []}

doc = Document()
doc.add_heading('DiasporicMemory: Bibliographie-Addendum', 0)
sub = doc.add_paragraph()
sub.add_run('Israelkorpus-Integration, 2026-06-15').italic = True

intro = doc.add_paragraph()
intro.add_run(
    f"Diese Datei dokumentiert die {len(review_items)} bibliographischen Neuzugaenge, "
    f"die im Rahmen der Israelkorpus-Integration am {today} in die Zotero-Gruppe "
    "DiasporicMemory aufgenommen wurden. Sie ergaenzt das vollstaendige Review-Dokument "
    "diasporicmemory_bib_2026-06-08.docx. Zusaetzlich wurde eine reine Metadaten-Registry "
    f"von {len(registry_items)} IS-Interviews (mit Handle-PIDs, ohne Transkripte) in der "
    "neuen Sammlung 'P9. Israelkorpus interviews (IS, registry)' angelegt; diese ist hier "
    "nicht im Detail aufgefuehrt."
)
doc.add_paragraph()

src = doc.add_paragraph()
src.add_run('Quelle der Neuzugaenge: ').bold = True
src.add_run(
    "Anne Bettens 'Sprache in der Emigration: Deutsch in Israel' (Stand April 2023); "
    "Flinz & Leonardi (2024), HTRes@LREC-COLING-2024; sowie die Zotero-Gruppe 'Israelkorpus' "
    "(Leonardi, 2219390) mit ihren Untersammlungen 'Publikationen von Interviewten' und "
    "'Publikationen zum IK'. Fuer DGD-Transkripte selbst wurden -- im Einklang mit den "
    "Nutzungsbedingungen der DGD, die das Bereitstellen der Daten an KI-Tools/Chatbots untersagen -- "
    "keine Volltext-Daten in das Projektrepository uebernommen."
)
doc.add_paragraph()

for ck in sorted(by_coll, key=lambda k: coll_path(k)):
    its = by_coll[ck]
    if not its: continue
    doc.add_heading(coll_path(ck), level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Eintrag'
    hdr[1].text = 'Notizen / Verifizierung'
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1.6)
    for it in sorted(its, key=lambda i: ((i['data'].get('date','') or '0000')[:4], i['data'].get('title',''))):
        d = it['data']
        row = table.add_row().cells
        row[0].text = cite(d)
        ts = tag_set(d)
        notes = []
        if 'Israelkorpus-Interviewter' in ts:
            notes.append('Primaerquelle (Verfasser*in ist Interviewpartner*in der IS/ISW/ISZ)')
        if 'Israelkorpus-Forschung' in ts:
            notes.append('Sekundaerliteratur zum Israelkorpus')
        row[1].text = '; '.join(notes)
    doc.add_paragraph()

doc.add_heading('P9. Israelkorpus interviews (IS, registry) -- Uebersicht', level=2)
doc.add_paragraph(
    f"{len(registry_items)} Interview-Metadatensaetze des Korpus IS (1. Generation) wurden als "
    "reine Registratur in die DiasporicMemory-Gruppe uebernommen. Jeder Eintrag enthaelt den "
    "Sprecher-Namen, das Interview-Datum, die Sigle (IS_E_NNNNN) und die Handle-PID. "
    "Sie dienen als Anker fuer kuenftige Zitationen und Quellenverweise und enthalten "
    "ausdruecklich KEINEN Transkript-Volltext. Sammlungsschluessel: 63X662X9."
)
doc.add_paragraph("Beispiel-Eintraege:")
for it in registry_items[:8]:
    d = it['data']
    line = f"  - {d.get('title','')}  ({d.get('url','')})"
    p = doc.add_paragraph(line)
    p.paragraph_format.left_indent = Inches(0.3)
doc.add_paragraph(f"  ... und {len(registry_items)-8} weitere.")

out = '/Users/sinairusinek/Documents/GitHub/JeckeArchive/data/bibliography/diasporicmemory_bib_addendum_2026-06-15.docx'
doc.save(out)
print(f"Saved: {out}  ({os.path.getsize(out)} bytes)")
